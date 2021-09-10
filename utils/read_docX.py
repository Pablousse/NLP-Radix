import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def get_full_text(filename):
    doc = docx.Document(filename)
    full_text = []

    # parsing links
    rels = doc.part.rels
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            full_text.append(rels[rel]._target)
    print('links found:', len(full_text))

    # parsing plain text paragraphs
    prgrphs = doc.paragraphs
    print('paragraphs found:', len(prgrphs))
    for paragraph in prgrphs:
        full_text.append(paragraph.text)

    # parsing text in tables
    tables = doc.tables
    print('tables found: ', len(tables))
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)
